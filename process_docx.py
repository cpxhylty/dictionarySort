from docx import Document
from docx.oxml.ns import qn
import re
import time

# steps:
# 1.源文件跑removeStartParaSpace，cmd+A cmd+shift+f9去除超链接
# 2.源文件最前边留一个空段落，跑modify宏
# 3.源文件英文字体全改为同中文
# 4.目标文件最后边需要有1234zzzz段落
# 5.删掉源文件的第一个空段落，跑process_docx
# 6.目标文件英文字体全改times

class Pinyin:
    def __init__(self, pinyin):
        self.content = pinyin  # real pinyin
        tableSort = str.maketrans('āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜü', 'aaaaooooeeeeiiiiuuuuuuuuu')
        self.contentSort = pinyin.translate(tableSort)  # pinyin for sort
        tableForSame = str.maketrans('āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜü', 'abcdabcdabcdabcdabcdabcda')
        self.contentForSame = pinyin.translate(tableForSame)  # pinyin for same letter sort

    def __lt__(self, other):
        if self.contentSort < other.contentSort:
            return True
        elif self.contentSort > other.contentSort:
            return False
        else:
            return self.contentForSame < other.contentForSame


def fromParatoPara(parato, parafrom):
    parato.paragraph_format.first_line_indent = parafrom.paragraph_format.first_line_indent
    parato.paragraph_format.left_indent = parafrom.paragraph_format.left_indent
    parato.paragraph_format.line_spacing = parafrom.paragraph_format.line_spacing
    for run in parafrom.runs:
        r = parato.add_run(run.text)
        if run.font.name is None:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            r.font.name = '宋体'
            r.font.name = run.font.name
            if r._element.rPr.rFonts is not None:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
        r.bold = run.bold

docxNamefrom = '金梦城1月（6551-6600） .docx'
docxNameto = 'test3.docx'

start = time.perf_counter()

document = Document(docxNamefrom)
document2 = Document(docxNameto)
startParas = []  # 每个词的第一段

# 把doc2已有词条加入列表(至少手动添加一个1234zzzz)
for paragraph in document2.paragraphs:
    if re.match('\d{4}', paragraph.text) is not None:
        startParas.append(paragraph)

count = 0
for paragraph in document.paragraphs:
    if re.match('\d{4}', paragraph.text) is not None:  # 是startpara
        count += 1
        print(paragraph.text[0:4])
        ParagraphPY = Pinyin(re.search('([āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜüa-z])+', paragraph.text).group(0))  # 生成Pinyin并比较,找插入的位置
        indexStartPara = 0
        for startPara in startParas:
            startParaPY = Pinyin(re.search('([āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜüa-z])+', startPara.text).group(0))
            if ParagraphPY < startParaPY:
                length = len(document2.paragraphs)
                for order in range(length):
                    if re.match('\d{4}', document2.paragraphs[order].text) is not None and document2.paragraphs[order].text == startPara.text:
                        paraBase = document2.paragraphs[order] # 在paraBase前插入新词条
                        break
                break
            indexStartPara += 1
        paraInsert = paraBase.insert_paragraph_before()
        fromParatoPara(paraInsert, paragraph)
        startParas.insert(indexStartPara, paragraph)
    else:  # 不是startpara
        paraInsert = paraBase.insert_paragraph_before()
        fromParatoPara(paraInsert, paragraph)
document2.save(docxNameto)
print('finish, time spent: {:.2f}seconds\nwords count: {}'.format(time.perf_counter()-start, count))