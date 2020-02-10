from docx import Document
from docx.oxml.ns import qn
import re
import time

def fromParatoPara(parato, parafrom):
    parato.paragraph_format.first_line_indent = parafrom.paragraph_format.first_line_indent
    parato.paragraph_format.left_indent = parafrom.paragraph_format.left_indent
    parato.paragraph_format.line_spacing = parafrom.paragraph_format.line_spacing
    for run in parafrom.runs:
        r = parato.add_run(run.text)
        if run.font.name is None:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            r.font.name = '宋体'
            r.font.name = run.font.name
            if r._element.rPr.rFonts is not None:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
        r.bold = run.bold

letter = 'M'
wordCount = 0
wordsPerFile = 20
fileCount = 1

docxName1 = 'test.docx'
docxName2 = letter+str(fileCount)+'.docx'

document1 = Document(docxName1)
document2 = Document(docxName2)


for paragraph in document1.paragraphs:
    if re.match('\d{4}', paragraph.text) is not None:
        wordCount += 1
        if wordCount == 21:
            wordCount = 1
            document2.save(docxName2)
            print('words in {}: 20'.format(docxName2))
            fileCount += 1
            docxName2 = letter + str(fileCount) + '.docx'
            document2 = Document(docxName2)
        paraInsert = document2.add_paragraph()
        fromParatoPara(paraInsert, paragraph)
    else:
        paraInsert = document2.add_paragraph()
        fromParatoPara(paraInsert, paragraph)

document2.save(docxName2)
print('words in {}: {}'.format(docxName2, wordCount))
print('finish, words count: {}'.format((fileCount-1) * wordsPerFile + wordCount))