from docx import Document
from docx.oxml.ns import qn
import re

docxName = input('input file name: ')

count = 0

document = Document(docxName)
length = len(document.paragraphs)
for index in range(length):
    if re.match('\d{4}', document.paragraphs[index].text) is not None:  # 是startpara
        count += 1
        print(document.paragraphs[index].text[0:4])
        document.paragraphs[index].text = re.sub('\s(?=[āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜüa-z])', '', document.paragraphs[index].text)
        for run in document.paragraphs[index].runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.name = 'Times New Roman'
document.save(docxName)
print('finish, words count: {}'.format(count))