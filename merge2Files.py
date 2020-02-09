from docx import Document
from docx.oxml.ns import qn
import re
import time

class Pinyin:
    def __init__(self, pinyin):
        self.content = pinyin  # real pinyin
        tableSort = str.maketrans('āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜü', 'aaaaooooeeeeiiiiuuuuuuuuu')
        self.contentSort = pinyin.translate(tableSort)  # pinyin for sort
        tableForSame = str.maketrans('āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜü', 'abcdabcdabcdabcdabcdabcda')
        self.contentForSame = pinyin.translate(tableForSame)  # pinyin for same letter sort

    def __lt__(self, other):
        if self.contentSort < other.contentSort:
            return True
        elif self.contentSort > other.contentSort:
            return False
        else:
            return self.contentForSame < other.contentForSame

def fromParatoPara(parato, parafrom):
    parato.paragraph_format.first_line_indent = parafrom.paragraph_format.first_line_indent
    parato.paragraph_format.left_indent = parafrom.paragraph_format.left_indent
    parato.paragraph_format.line_spacing = parafrom.paragraph_format.line_spacing
    for run in parafrom.runs:
        r = parato.add_run(run.text)
        if run.font.name is None:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            r.font.name = '宋体'
            r.font.name = run.font.name
            if r._element.rPr.rFonts is not None:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
        r.bold = run.bold

start = time.perf_counter()

docxName1 = '金梦城1月（6551-6600）.docx'
docxName2 = '李安涛1月（6601-6650）.docx'
docxtoName = 'test.docx'

document1 = Document(docxName1)
document2 = Document(docxName2)
documentto = Document(docxtoName)

length1 = len(document1.paragraphs)
length2 = len(document2.paragraphs)
index1 = 0
index2 = 0

state = 0

while True:
    pinyin1 = Pinyin(re.search('([āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜüa-z])+', document1.paragraphs[index1].text).group(0))
    pinyin2 = Pinyin(re.search('([āáǎàōóǒòēéěèīíǐìūúǔùǖǘǚǜüa-z])+', document2.paragraphs[index2].text).group(0))
    if pinyin1 < pinyin2:
        while True:
            paraInsert = documentto.add_paragraph()
            fromParatoPara(paraInsert, document1.paragraphs[index1])
            index1 += 1
            if index1 == length1:
                break
            if re.match('\d{4}', document1.paragraphs[index1].text) is not None:
                break
    else:
        while True:
            paraInsert = documentto.add_paragraph()
            fromParatoPara(paraInsert, document2.paragraphs[index2])
            index2 += 1
            if index2 == length2:
                break
            if re.match('\d{4}', document2.paragraphs[index2].text) is not None:
                break
    if index1 == length1:
        state = 2
        break
    if index2 == length2:
        state = 1
        break

if state == 1:
    for index in range(index1, length1):
        paraInsert = documentto.add_paragraph()
        fromParatoPara(paraInsert, document1.paragraphs[index])
else:
    for index in range(index2, length2):
        paraInsert = documentto.add_paragraph()
        fromParatoPara(paraInsert, document2.paragraphs[index])

documentto.save(docxtoName)
print('finish, time spent: {:.2f}seconds'.format(time.perf_counter()-start))